"""
Offerte docx-extractie service.

Doel: een geüpload .docx bestand omzetten naar leesbare ruwe tekst,
met correcte afhandeling van samengevoegde tabelcellen (het patroon
dat we bij de Brink-offerte tegenkwamen: een hele tabel die eigenlijk
één grote merged cell is met spatie-uitlijning).

Dit vervangt GEEN semantische parsing (welk getal hoort bij welk item) -
dat blijft de taak van de Claude-call verderop in de n8n flow. Deze
service zorgt alleen dat er schone, correcte tekst uit het bestand komt
om aan die Claude-call te voeren, zonder dat python-docx eigenaardigheden
(gedupliceerde merged cells, kapotte tabellen) de output vervuilen.

Draai lokaal:
    pip install -r requirements.txt
    uvicorn main:app --host 0.0.0.0 --port 8000

Endpoint:
    POST /extract-docx
    multipart/form-data, veld "file" = het .docx bestand

    Response:
    {
      "raw_text": "... alle uitgelezen tekst ...",
      "tables_found": 2,
      "tables_used": 1,
      "warnings": ["Tabel 2 was volledig leeg, overgeslagen."]
    }
"""

import io
import zipfile
import base64
from fastapi import FastAPI, UploadFile, File, HTTPException
from docx import Document

try:
    from PIL import Image
except ImportError:
    Image = None

app = FastAPI(title="Offerte docx-extractie service")

MIN_IMAGE_DIMENSION = 60      # kleinere afbeeldingen (bullets, iconen) worden overgeslagen
THUMBNAIL_MAX_WIDTH = 400     # basis64-thumbnails blijven klein, geen volledige resolutie nodig voor een keuzelijst
MAX_IMAGES_RETURNED = 8


def extract_embedded_images(file_bytes):
    """Haal alle ingebedde afbeeldingen uit een docx (via word/media/), zodat de
    gebruiker er zelf een logo uit kan kiezen in plaats van dat het systeem raadt
    welke afbeelding de logo is. Slaat te kleine afbeeldingen (iconen, bullets) over."""
    images = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            media_files = sorted(n for n in z.namelist() if n.startswith("word/media/"))
            for name in media_files:
                if len(images) >= MAX_IMAGES_RETURNED:
                    break
                raw = z.read(name)
                if Image is None:
                    continue
                try:
                    img = Image.open(io.BytesIO(raw))
                    w, h = img.size
                except Exception:
                    continue
                if w < MIN_IMAGE_DIMENSION or h < MIN_IMAGE_DIMENSION:
                    continue
                if w > THUMBNAIL_MAX_WIDTH:
                    ratio = THUMBNAIL_MAX_WIDTH / w
                    img = img.convert("RGB").resize((THUMBNAIL_MAX_WIDTH, int(h * ratio)))
                buf = io.BytesIO()
                img.convert("RGB").save(buf, format="JPEG", quality=80)
                images.append({
                    "filename": name.rsplit("/", 1)[-1],
                    "width": w,
                    "height": h,
                    "thumbnail_base64": base64.b64encode(buf.getvalue()).decode(),
                })
    except zipfile.BadZipFile:
        pass
    return images


def dedupe_row_cells(row):
    """Geef elke onderliggende tabelcel van een rij precies één keer terug,
    ook als de rij horizontaal samengevoegde cellen bevat (die anders
    meerdere keren achter elkaar worden geretourneerd door python-docx)."""
    seen = set()
    cells = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid in seen:
            continue
        seen.add(cid)
        cells.append(cell)
    return cells


def extract_table_text(table):
    """Zet een hele tabel om naar leesbare tekst, rij voor rij, met de
    cellen van elke rij gescheiden door ' | '. Slaat volledig lege rijen
    over (zoals de 37 lege rijen die we bij Brink tegenkwamen)."""
    lines = []
    for row in table.rows:
        cells = dedupe_row_cells(row)
        cell_texts = [c.text.strip() for c in cells]
        if not any(cell_texts):
            continue
        lines.append(" | ".join(cell_texts))
    return "\n".join(lines)


@app.post("/extract-docx")
async def extract_docx(file: UploadFile = File(...)):
    if not file.filename.lower().endswith((".docx",)):
        raise HTTPException(status_code=400, detail="Alleen .docx bestanden worden ondersteund door dit endpoint.")

    content = await file.read()
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Kon het bestand niet openen als docx: {e}")

    warnings = []
    used_blocks = []

    # Tekst buiten tabellen (voorblad-teksten, losse notities e.d.)
    body_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    if body_text:
        used_blocks.append("TEKST BUITEN TABELLEN:\n" + body_text)

    tables_used = 0
    for i, table in enumerate(doc.tables):
        table_text = extract_table_text(table)
        if not table_text.strip():
            warnings.append(f"Tabel {i + 1} was volledig leeg, overgeslagen.")
            continue
        used_blocks.append(f"TABEL {i + 1}:\n" + table_text)
        tables_used += 1

    raw_text = "\n\n---\n\n".join(used_blocks)

    if not raw_text.strip():
        warnings.append("Geen bruikbare tekst gevonden in het hele document.")

    images = extract_embedded_images(content)
    if not images:
        warnings.append("Geen (bruikbare) afbeeldingen gevonden om als logo te kiezen.")

    return {
        "raw_text": raw_text,
        "tables_found": len(doc.tables),
        "tables_used": tables_used,
        "images": images,
        "warnings": warnings,
    }


@app.get("/health")
async def health():
    return {"status": "ok"}


# ---------------------------------------------------------------------------
# Generatie: JSON (output van de Claude-parse-stap) -> Word-document
# ---------------------------------------------------------------------------

import io as _io
import os
import base64
from typing import Optional, List
from pydantic import BaseModel, model_validator, field_validator
from fastapi.responses import StreamingResponse
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

FONT_NAME = "Calibri"
HEADER_BG = "2F5233"
SHADE_BG = "F2F2F2"
TOTAL_BG = "E4E4E4"
REFURNITY_GREEN = "8DC63F"
REFURNITY_LOGO_PATH = os.path.join(os.path.dirname(__file__), "refurnity_logo.png")
REFURNITY_ADDRESS_LINES = [
    "Refurnity BV",
    "Veerstraat 47 II, 1075 SN Amsterdam",
    "Telefoon 0621502536",
    "Mail info@ReFurnity.nl",
    "www.ReFurnity.nl",
]
REFURNITY_NAAM = "ReFurnity BV"
REFURNITY_ADRES_KORT = ["Veerstraat 47 II", "1075 SN Amsterdam"]
REFURNITY_BANK = "NL28 RABO 0146 2478 09"
REFURNITY_BTW = "NL8176.80.056.B01"
REFURNITY_KVK = "12065660"


def _leeg_naar_none(v):
    """Lovable/n8n sturen lege getalvelden soms als lege string ("") in plaats
    van null. Pydantic accepteert dat niet als geldig getal en gooit dan een
    422 ("Input should be a valid number"). Vang dat hier af voor validatie."""
    if isinstance(v, str) and v.strip() == "":
        return None
    return v


class Regel(BaseModel):
    item: str
    omschrijving: Optional[str] = ""
    specs: Optional[List[str]] = []
    aantal: Optional[float] = None
    prijs_per_stuk: Optional[float] = None
    totaal: float
    waarschuwing: Optional[str] = None
    opmerking: Optional[str] = None
    foto_base64: Optional[str] = None  # optionele productfoto, ruwe base64 zonder data:-prefix

    @field_validator("aantal", "prijs_per_stuk", mode="before")
    @classmethod
    def _leeg_naar_none_of_nul(cls, v):
        return _leeg_naar_none(v)

    @field_validator("totaal", mode="before")
    @classmethod
    def _totaal_leeg_naar_nul(cls, v):
        # "totaal" is verplicht in het model, maar een gebruiker kan een
        # nieuwe regel toevoegen en het bedrag nog even leeg laten staan.
        # Een lege string mag dan niet de hele generatie laten crashen.
        v = _leeg_naar_none(v)
        return 0.0 if v is None else v

    @model_validator(mode="after")
    def _vul_prijs_per_stuk_aan(self):
        # Als iemand alleen aantal + totaal invult (bijv. handmatig een regel
        # toevoegen in het dashboard) en de prijs per stuk leeg laat, reken
        # 'm hier uit zodat de kolom nooit onnodig leeg blijft.
        if self.prijs_per_stuk is None and self.aantal:
            self.prijs_per_stuk = round(self.totaal / self.aantal, 2)
        return self


class Klant(BaseModel):
    naam: Optional[str] = None
    adres: Optional[str] = None
    contact: Optional[str] = None
    logo_base64: Optional[str] = None  # ruwe base64, zonder data:-prefix


class AanbetalingTermijn(BaseModel):
    label: str
    percentage: float  # bijv. 50 voor 50%

    @field_validator("percentage", mode="before")
    @classmethod
    def _leeg_naar_nul(cls, v):
        v = _leeg_naar_none(v)
        return 0.0 if v is None else v


class DocumentGegevens(BaseModel):
    # Orderbevestiging
    ordernummer: Optional[str] = None
    projectmanager: Optional[str] = None
    accountmanager: Optional[str] = None
    levertermijn: Optional[str] = None
    betalingstermijn: Optional[str] = None
    # Factuur
    factuurnummer: Optional[str] = None
    factuurdatum: Optional[str] = None
    vervaldatum: Optional[str] = None
    # Gedeeld
    datum: Optional[str] = None
    referentie: Optional[str] = None  # "Uw referentie" / kenmerk
    gebaseerd_op: Optional[str] = None  # verwijzing naar offerte- of ordernummer


class GenerateRequest(BaseModel):
    regels: List[Regel]
    algemene_opmerkingen: Optional[List[str]] = []
    klant: Optional[Klant] = None
    template: Optional[str] = "1"
    toeslag_percentage: Optional[float] = None  # bijv. 20 voor 20% opslag op het leveranciersbedrag
    korting_percentage: Optional[float] = None  # bijv. 10 voor 10% korting, altijd zichtbaar voor de klant
    btw_percentage: Optional[float] = 21  # standaard 21%, zet expliciet op 0 om BTW-regel weg te laten
    aanbetaling_termijnen: Optional[List[AanbetalingTermijn]] = None
    document_type: Optional[str] = "offerte"  # "offerte" | "orderbevestiging" | "factuur"
    document: Optional[DocumentGegevens] = None

    @field_validator("toeslag_percentage", mode="before")
    @classmethod
    def _toeslag_leeg_naar_none(cls, v):
        return _leeg_naar_none(v)

    @field_validator("korting_percentage", mode="before")
    @classmethod
    def _korting_leeg_naar_none(cls, v):
        return _leeg_naar_none(v)

    @field_validator("btw_percentage", mode="before")
    @classmethod
    def _btw_leeg_naar_default(cls, v):
        v = _leeg_naar_none(v)
        return 21 if v is None else v



def _shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


def _set_font(run, size=10, bold=False, color=None, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _fix_openpyxl_lege_formule_waarde(xlsx_bytes: bytes) -> bytes:
    """Repareert twee losstaande openpyxl-eigenaardigheden (bevestigd aanwezig
    in een kale, lege openpyxl 3.1.5-workbook, dus geen fout in onze eigen
    code) die Microsoft Excel laten weigeren met "bestand is beschadigd",
    terwijl lenient tools als openpyxl zelf en LibreOffice ze gewoon
    negeren:

    1. Formule-cellen krijgen een lege `<v></v>` mee (bijv.
       `<f>C6*D6</f><v></v>`). Excel wil hier ofwel een geldige gecachete
       waarde, of helemaal geen <v>-tag.
    2. Relatiebestanden (*.rels) gebruiken soms een absoluut pad
       (bijv. Target="/xl/drawings/drawing1.xml") terwijl andere relaties
       in datzelfde bestand relatieve paden gebruiken. Die inconsistentie
       lost Excel niet netjes op. We maken elk absoluut pad relatief ten
       opzichte van de map van het onderdeel waar het .rels-bestand bij
       hoort, ongeacht hoeveel van zulke relaties er zijn (dus dit werkt
       ook bij meerdere foto's, extra werkbladen, etc.).
    """
    import re as _re
    import posixpath as _pp

    def _fix_rels(rels_filename: str, data: bytes) -> bytes:
        if rels_filename == "_rels/.rels":
            parent_dir = ""
        else:
            parent_dir = rels_filename.split("/_rels/", 1)[0]

        def _repl(m):
            target = m.group(1).decode()
            if not target.startswith("/"):
                return m.group(0)
            rel = _pp.relpath(target.lstrip("/"), parent_dir) if parent_dir else target.lstrip("/")
            return b'Target="' + rel.encode() + b'"'

        return _re.sub(rb'Target="([^"]+)"', _repl, data)

    in_buf = _io.BytesIO(xlsx_bytes)
    out_buf = _io.BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("xl/worksheets/") and item.filename.endswith(".xml") and "_rels" not in item.filename:
                data = _re.sub(rb"(</f>)<v></v>", rb"\1", data)
            elif item.filename.endswith(".rels"):
                data = _fix_rels(item.filename, data)
            zout.writestr(item, data)
    return out_buf.getvalue()


def _money(n):
    return f"€ {n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _verwerk_foto_bytes(raw_bytes: bytes, item_naam: str = "?", max_breedte_px: int = 1000) -> bytes:
    """Normaliseert een geuploade productfoto voor het inladen in Word/Excel.

    Vangt de twee meest voorkomende faalredenen af die eerder stil werden
    weggeslikt (bare except): CMYK/palette-kleurmodus (waar python-docx en
    openpyxl niet altijd mee overweg kunnen) en onnodig grote pixelafmetingen
    (die de foto in het document ook nodeloos groot en traag maken, terwijl
    hij daar toch maar een paar centimeter breed wordt getoond).
    """
    img = PILImage.open(_io.BytesIO(raw_bytes))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    if img.width > max_breedte_px:
        nieuwe_hoogte = round(img.height * (max_breedte_px / img.width))
        img = img.resize((max_breedte_px, nieuwe_hoogte), PILImage.LANCZOS)
    buf = _io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue()


def _tbl_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _fill_info_cell(cell, pairs):
    cell.paragraphs[0].text = ""
    first = True
    for label, value in pairs:
        if not value:
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(label + " ")
        _set_font(run, size=10, bold=True)
        run2 = p.add_run(str(value))
        _set_font(run2, size=10)
    if first:
        cell.paragraphs[0].add_run("")


def _build_item_table_and_totals(doc, data: GenerateRequest, toon_details: bool = True) -> float:
    """Bouwt de regel-tabel + totalenblok + aanbetalingstermijnen + opmerkingen.
    Gedeeld door offerte, orderbevestiging en factuur zodat ze niet uit elkaar
    kunnen groeien. Retourneert het eindtotaal (incl. eventuele BTW/toeslag).

    toon_details=False (orderbevestiging/factuur): alleen item, omschrijving als
    platte tekst, aantal, prijs, totaal. Geen specs-bullets en geen productfoto's,
    dat is voor een klant te veel ruis op een order-/betaaldocument. De offerte
    blijft wel het volledige detail tonen (toon_details=True)."""
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [Cm(3.5), Cm(6), Cm(1.8), Cm(2.5), Cm(2.7)]
    headers = ["Item", "Omschrijving / specificaties", "Aantal", "Prijs p.s.", "Totaal"]

    hdr_cells = table.rows[0].cells
    for i, (htext, w) in enumerate(zip(headers, widths)):
        hdr_cells[i].width = w
        _shade_cell(hdr_cells[i], HEADER_BG)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(htext)
        _set_font(run, size=10, bold=True, color="FFFFFF")
    _repeat_header(table.rows[0])

    grand_total = 0.0
    for idx, regel in enumerate(data.regels):
        row = table.add_row()
        shade = idx % 2 == 1
        for i, w in enumerate(widths):
            row.cells[i].width = w
            if shade:
                _shade_cell(row.cells[i], SHADE_BG)

        p = row.cells[0].paragraphs[0]
        run = p.add_run(regel.item)
        _set_font(run, bold=True)

        if toon_details and regel.foto_base64:
            try:
                foto_bytes = base64.b64decode(regel.foto_base64)
                foto_bytes = _verwerk_foto_bytes(foto_bytes, item_naam=regel.item)
                fp = row.cells[0].add_paragraph()
                frun = fp.add_run()
                frun.add_picture(_io.BytesIO(foto_bytes), width=Cm(2.8))
            except Exception as e:
                print(f"[foto-fout] Kon foto voor '{regel.item}' niet in het Word-document plaatsen: {type(e).__name__}: {e}")

        cell1 = row.cells[1]
        cell1.paragraphs[0].text = ""
        first = True
        if regel.omschrijving:
            par = cell1.paragraphs[0] if first else cell1.add_paragraph()
            run = par.add_run(regel.omschrijving)
            _set_font(run)
            first = False
        if toon_details:
            for spec in (regel.specs or []):
                par = cell1.paragraphs[0] if first else cell1.add_paragraph()
                run = par.add_run(f"•  {spec}")
                _set_font(run, size=9, color="444444")
                first = False
            if regel.opmerking:
                par = cell1.paragraphs[0] if first else cell1.add_paragraph()
                run = par.add_run(regel.opmerking)
                _set_font(run, size=9, color="555555")
                run.italic = True
                first = False
        if first:
            cell1.paragraphs[0].add_run("")

        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(str(int(regel.aantal)) if regel.aantal is not None else "")
        _set_font(run)

        p = row.cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(_money(regel.prijs_per_stuk) if regel.prijs_per_stuk is not None else "")
        _set_font(run)

        p = row.cells[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(_money(regel.totaal))
        _set_font(run, bold=True)

        grand_total += regel.totaal

    def _add_summary_row(label, bedrag, bold=True, shade_bg=None):
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        if shade_bg:
            _shade_cell(row.cells[0], shade_bg)
            _shade_cell(row.cells[4], shade_bg)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        _set_font(run, size=11 if bold else 10, bold=bold)
        p = row.cells[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(_money(bedrag))
        _set_font(run, size=11 if bold else 10, bold=bold)

    subtotaal = grand_total
    toeslag_percentage = data.toeslag_percentage or 0
    toeslag_bedrag = subtotaal * (toeslag_percentage / 100) if toeslag_percentage else 0.0
    na_toeslag = subtotaal + toeslag_bedrag

    korting_percentage = data.korting_percentage or 0
    korting_bedrag = na_toeslag * (korting_percentage / 100) if korting_percentage else 0.0
    totaal_excl_btw = na_toeslag - korting_bedrag

    heeft_btw = bool(data.btw_percentage)
    toon_toeslag_breakdown = toon_details and bool(toeslag_bedrag)
    # Korting wordt, net als Toeslag, nergens als aparte regel getoond -
    # alleen stilzwijgend verwerkt in het eindtotaal.
    heeft_enige_breakdown = toon_toeslag_breakdown

    if not heeft_enige_breakdown and not heeft_btw:
        _add_summary_row("TOTAAL", totaal_excl_btw, shade_bg=TOTAL_BG)
        eindtotaal = totaal_excl_btw
    else:
        if toon_toeslag_breakdown:
            _add_summary_row("Subtotaal", subtotaal, bold=False)
            _add_summary_row(f"Toeslag ({toeslag_percentage:g}%)", toeslag_bedrag, bold=False)

        if heeft_btw:
            _add_summary_row("Totaal excl. BTW", totaal_excl_btw, bold=False)
            btw_bedrag = totaal_excl_btw * (data.btw_percentage / 100)
            _add_summary_row(f"BTW ({data.btw_percentage:g}%)", btw_bedrag, bold=False)
            eindtotaal = totaal_excl_btw + btw_bedrag
            _add_summary_row("Totaal incl. BTW", eindtotaal, shade_bg=TOTAL_BG)
        else:
            eindtotaal = totaal_excl_btw
            _add_summary_row("TOTAAL", eindtotaal, shade_bg=TOTAL_BG)

    if toon_details and data.aanbetaling_termijnen:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Betaaltermijnen")
        _set_font(run, size=11, bold=True)
        termijnen = data.aanbetaling_termijnen
        lopend_totaal = 0.0
        for idx, termijn in enumerate(termijnen):
            if idx == len(termijnen) - 1:
                bedrag = round(eindtotaal - lopend_totaal, 2)
            else:
                bedrag = round(eindtotaal * (termijn.percentage / 100), 2)
                lopend_totaal += bedrag
            p = doc.add_paragraph()
            run = p.add_run(f"{termijn.label}: {termijn.percentage:g}% = {_money(bedrag)}")
            _set_font(run, size=10)

    # Let op: algemene_opmerkingen worden bewust NIET in het gegenereerde
    # document geschreven. Dit veld blijft alleen zichtbaar in Lovable zelf
    # (via de API-respons van de parse-stap), niet in wat de klant ontvangt.

    return eindtotaal


def _build_offerte(doc, data: GenerateRequest):
    section = doc.sections[0]

    section.different_first_page_header_footer = True
    if os.path.exists(REFURNITY_LOGO_PATH):
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run()
        hrun.add_picture(REFURNITY_LOGO_PATH, width=Cm(2.8))
        section.first_page_header.paragraphs[0].text = ""

    if data.klant and data.klant.logo_base64:
        try:
            logo_bytes = base64.b64decode(data.klant.logo_base64)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(_io.BytesIO(logo_bytes), width=Cm(6))
        except Exception:
            pass

    if data.klant and data.klant.naam:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Offerte voor {data.klant.naam}")
        _set_font(run, size=18, bold=True)
        if data.klant.adres:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(data.klant.adres)
            _set_font(run2, size=10)
        if data.klant.contact:
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run(data.klant.contact)
            _set_font(run3, size=10)

    doc.add_paragraph()
    if os.path.exists(REFURNITY_LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(REFURNITY_LOGO_PATH, width=Cm(7))
    for line in REFURNITY_ADDRESS_LINES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        _set_font(run, size=9, color="666666")

    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Offerte")
    _set_font(run, size=18, bold=True)

    _build_item_table_and_totals(doc, data)


def _build_orderbevestiging(doc, data: GenerateRequest):
    section = doc.sections[0]
    if os.path.exists(REFURNITY_LOGO_PATH):
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run()
        hrun.add_picture(REFURNITY_LOGO_PATH, width=Cm(2.8))

    d = data.document or DocumentGegevens()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Orderbevestiging")
    _set_font(run, size=20, bold=True, color=HEADER_BG)

    if data.klant and data.klant.naam:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(data.klant.naam)
        _set_font(run, size=13, bold=True)
        if data.klant.adres or data.klant.contact:
            sub2 = doc.add_paragraph()
            sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tekst = "  ·  ".join(x for x in [data.klant.adres, data.klant.contact] if x)
            run = sub2.add_run(tekst)
            _set_font(run, size=10, color="555555")

    doc.add_paragraph()

    outer = doc.add_table(rows=1, cols=2)
    outer.autofit = False
    outer.columns[0].width = Cm(8)
    outer.columns[1].width = Cm(8)
    left_cell, right_cell = outer.rows[0].cells
    _fill_info_cell(left_cell, [
        ("Ordernummer:", d.ordernummer),
        ("Datum:", d.datum),
        ("Uw referentie:", d.referentie),
    ])
    _fill_info_cell(right_cell, [
        ("Projectmanager:", d.projectmanager),
        ("Accountmanager:", d.accountmanager),
        ("Gebaseerd op offerte:", d.gebaseerd_op),
    ])
    _tbl_no_borders(outer)

    doc.add_paragraph()
    _build_item_table_and_totals(doc, data, toon_details=False)

    doc.add_paragraph()
    for label, value in [("Levertermijn:", d.levertermijn), ("Betalingstermijn:", d.betalingstermijn)]:
        if not value:
            continue
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        _set_font(run, size=10, bold=True)
        run2 = p.add_run(value)
        _set_font(run2, size=10)

    p = doc.add_paragraph()
    run = p.add_run("Hartelijk dank voor uw bestelling.")
    _set_font(run, size=10, italic=True, color=REFURNITY_GREEN)


def _build_factuur(doc, data: GenerateRequest):
    d = data.document or DocumentGegevens()

    outer = doc.add_table(rows=1, cols=2)
    outer.autofit = False
    outer.columns[0].width = Cm(6)
    outer.columns[1].width = Cm(10)
    logo_cell, info_cell = outer.rows[0].cells

    logo_cell.paragraphs[0].text = ""
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if os.path.exists(REFURNITY_LOGO_PATH):
        p = logo_cell.paragraphs[0]
        run = p.add_run()
        run.add_picture(REFURNITY_LOGO_PATH, width=Cm(5.5))

    info_cell.paragraphs[0].text = ""
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    info_lines = [REFURNITY_NAAM] + REFURNITY_ADRES_KORT + [
        f"Bank: {REFURNITY_BANK}",
        f"BTW: {REFURNITY_BTW}",
        f"KVK: {REFURNITY_KVK}",
    ]
    for i, line in enumerate(info_lines):
        p = info_cell.paragraphs[0] if i == 0 else info_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        _set_font(run, size=9, bold=(i == 0), color="333333" if i == 0 else "666666")
    _tbl_no_borders(outer)

    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(8)
    hr.paragraph_format.space_after = Pt(14)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), REFURNITY_GREEN)
    pBdr.append(bottom)
    pPr.append(pBdr)

    if data.klant and data.klant.naam:
        client_table = doc.add_table(rows=1, cols=1)
        client_table.autofit = False
        client_table.columns[0].width = Cm(16)
        c = client_table.rows[0].cells[0]
        c.paragraphs[0].text = ""
        lines = [data.klant.naam]
        if data.klant.contact:
            lines.append(f"T.a.v. {data.klant.contact}")
        if data.klant.adres:
            lines.append(data.klant.adres)
        for i, line in enumerate(lines):
            p = c.paragraphs[0] if i == 0 else c.add_paragraph()
            run = p.add_run(line)
            _set_font(run, size=10, bold=(i == 0))
        _tbl_no_borders(client_table)

    doc.add_paragraph()

    title = doc.add_paragraph()
    run = title.add_run(f"Factuur {d.factuurnummer or ''}".strip())
    _set_font(run, size=20, bold=True, color=HEADER_BG)

    if d.referentie or d.gebaseerd_op:
        p = doc.add_paragraph()
        run = p.add_run("Kenmerk: ")
        _set_font(run, size=10, bold=True)
        tekst = "  ·  ".join(x for x in [d.referentie, f"Gebaseerd op {d.gebaseerd_op}" if d.gebaseerd_op else None] if x)
        run2 = p.add_run(tekst)
        _set_font(run2, size=10)

    outer2 = doc.add_table(rows=1, cols=2)
    outer2.autofit = False
    outer2.columns[0].width = Cm(4)
    outer2.columns[1].width = Cm(4)
    c1, c2 = outer2.rows[0].cells
    _fill_info_cell(c1, [("Factuurdatum:", d.factuurdatum)])
    _fill_info_cell(c2, [("Vervaldatum:", d.vervaldatum)])
    _tbl_no_borders(outer2)

    doc.add_paragraph()
    eindtotaal = _build_item_table_and_totals(doc, data, toon_details=False)

    doc.add_paragraph()
    p = doc.add_paragraph()
    vervaldatum_tekst = f"vóór {d.vervaldatum} " if d.vervaldatum else ""
    factuurnr_tekst = f", onder vermelding van {d.factuurnummer}" if d.factuurnummer else ""
    run = p.add_run(
        f"Wij verzoeken u vriendelijk het bovenstaande bedrag van {_money(eindtotaal)} {vervaldatum_tekst}"
        f"over te maken naar bovenstaand rekeningnummer{factuurnr_tekst}. "
        "Voor vragen kunt u contact opnemen via info@ReFurnity.nl."
    )
    _set_font(run, size=9, color="555555")


def build_docx_template_1(data: GenerateRequest) -> bytes:
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    doc_type = data.document_type or "offerte"
    if doc_type == "orderbevestiging":
        _build_orderbevestiging(doc, data)
    elif doc_type == "factuur":
        _build_factuur(doc, data)
    else:
        _build_offerte(doc, data)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

TEMPLATE_BUILDERS = {
    "1": build_docx_template_1,
    "2": build_docx_template_1,  # TODO: eigen compacte variant
    "3": build_docx_template_1,  # TODO: eigen huisstijl-variant
}


@app.post("/generate-docx")
async def generate_docx(data: GenerateRequest):
    builder = TEMPLATE_BUILDERS.get(data.template or "1")
    if builder is None:
        raise HTTPException(status_code=400, detail=f"Onbekend template: {data.template}")

    docx_bytes = builder(data)
    doc_type = data.document_type or "offerte"
    filename = f"{doc_type}.docx"
    return StreamingResponse(
        _io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


import subprocess
import tempfile
import uuid


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Converteer docx-bytes naar pdf-bytes via headless LibreOffice.
    Elke aanroep gebruikt een eigen tijdelijke map, zodat gelijktijdige
    requests elkaar niet in de weg zitten."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, f"{uuid.uuid4().hex}.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [
                "soffice", "--headless", "--norestore",
                "--convert-to", "pdf", "--outdir", tmpdir, docx_path,
            ],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice-conversie mislukt: {result.stderr}")

        pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"Geen PDF geproduceerd. stdout: {result.stdout}, stderr: {result.stderr}")

        with open(pdf_path, "rb") as f:
            return f.read()


@app.post("/generate-pdf")
async def generate_pdf(data: GenerateRequest):
    builder = TEMPLATE_BUILDERS.get(data.template or "1")
    if builder is None:
        raise HTTPException(status_code=400, detail=f"Onbekend template: {data.template}")

    docx_bytes = builder(data)
    try:
        pdf_bytes = convert_docx_to_pdf(docx_bytes)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    doc_type = data.document_type or "offerte"
    return StreamingResponse(
        _io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{doc_type}.pdf"'},
    )


# ---------------------------------------------------------------------------
# Generatie: JSON -> Excel-document, met echte formules (geen losse
# Python-berekeningen), zodat het naspeurbaar en herrekenbaar blijft in Excel.
# ---------------------------------------------------------------------------

from openpyxl import Workbook
from openpyxl.styles import Font as XFont, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage

XFONT_NAME = "Calibri"
XHEADER_FILL = "2F5233"
XTOTAL_FILL = "E4E4E4"


def build_xlsx(data: GenerateRequest) -> bytes:
    wb = Workbook()
    ws = wb.active
    doc_type = data.document_type or "offerte"
    doc_titel = {"orderbevestiging": "Orderbevestiging", "factuur": "Factuur"}.get(doc_type, "Offerte")
    ws.title = doc_titel
    d = data.document or DocumentGegevens()

    col_widths = {"A": 30, "B": 40, "C": 10, "D": 14, "E": 14}
    for col, w in col_widths.items():
        ws.column_dimensions[col].width = w

    row = 1

    if data.klant and data.klant.logo_base64:
        try:
            logo_bytes = base64.b64decode(data.klant.logo_base64)
            logo_bytes = _verwerk_foto_bytes(logo_bytes, item_naam="klantlogo", max_breedte_px=400)
            pil_logo = PILImage.open(_io.BytesIO(logo_bytes))
            schaal = min(1.0, 150 / pil_logo.width)
            img = XLImage(_io.BytesIO(logo_bytes))
            img.width = round(pil_logo.width * schaal)
            img.height = round(pil_logo.height * schaal)
            ws.add_image(img, f"A{row}")
            ws.row_dimensions[row].height = max(60, img.height * 0.8)
            row += 2
        except Exception as e:
            print(f"[foto-fout] Kon klantlogo niet in het Excel-bestand plaatsen: {type(e).__name__}: {e}")

    if data.klant and data.klant.naam:
        ws.merge_cells(f"A{row}:E{row}")
        c = ws.cell(row=row, column=1, value=f"{doc_titel} voor {data.klant.naam}")
        c.font = XFont(name=XFONT_NAME, size=14, bold=True)
        c.alignment = Alignment(horizontal="center")
        row += 1
        if data.klant.adres:
            ws.merge_cells(f"A{row}:E{row}")
            c = ws.cell(row=row, column=1, value=data.klant.adres)
            c.font = XFont(name=XFONT_NAME, size=10)
            c.alignment = Alignment(horizontal="center")
            row += 1
        if data.klant.contact:
            ws.merge_cells(f"A{row}:E{row}")
            c = ws.cell(row=row, column=1, value=data.klant.contact)
            c.font = XFont(name=XFONT_NAME, size=10)
            c.alignment = Alignment(horizontal="center")
            row += 1
    row += 1

    # Documentgegevens (ordernummer/datum/etc.) - zelfde velden per documenttype
    # als de Word-versie, alleen platte "Label: waarde"-regels i.p.v. een tabel.
    if doc_type == "orderbevestiging":
        info_pairs = [
            ("Ordernummer", d.ordernummer), ("Datum", d.datum), ("Uw referentie", d.referentie),
            ("Projectmanager", d.projectmanager), ("Accountmanager", d.accountmanager),
            ("Gebaseerd op offerte", d.gebaseerd_op),
        ]
    elif doc_type == "factuur":
        info_pairs = [
            ("Factuurnummer", d.factuurnummer), ("Factuurdatum", d.factuurdatum),
            ("Vervaldatum", d.vervaldatum), ("Uw referentie", d.referentie),
            ("Gebaseerd op", d.gebaseerd_op),
        ]
    else:
        info_pairs = [("Datum", d.datum)]

    for label, waarde in info_pairs:
        if not waarde:
            continue
        c = ws.cell(row=row, column=1, value=f"{label}:")
        c.font = XFont(name=XFONT_NAME, size=10, bold=True)
        ws.cell(row=row, column=2, value=str(waarde)).font = XFont(name=XFONT_NAME, size=10)
        row += 1
    if any(w for _, w in info_pairs):
        row += 1

    header_row = row
    headers = ["Item", "Omschrijving / specificaties", "Aantal", "Prijs p.s.", "Totaal"]
    for i, htext in enumerate(headers, start=1):
        c = ws.cell(row=header_row, column=i, value=htext)
        c.font = XFont(name=XFONT_NAME, size=10, bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=XHEADER_FILL)
    row += 1

    first_item_row = row
    for regel in data.regels:
        omschrijving_delen = []
        if regel.omschrijving:
            omschrijving_delen.append(regel.omschrijving)
        for spec in (regel.specs or []):
            omschrijving_delen.append(f"• {spec}")
        if regel.opmerking:
            omschrijving_delen.append(regel.opmerking)

        ws.cell(row=row, column=1, value=regel.item).font = XFont(name=XFONT_NAME, size=10, bold=True)
        c = ws.cell(row=row, column=2, value="\n".join(omschrijving_delen))
        c.font = XFont(name=XFONT_NAME, size=10)
        c.alignment = Alignment(wrap_text=True, vertical="top")

        if regel.aantal is not None:
            ws.cell(row=row, column=3, value=regel.aantal).font = XFont(name=XFONT_NAME, size=10)
        if regel.prijs_per_stuk is not None:
            pc = ws.cell(row=row, column=4, value=regel.prijs_per_stuk)
            pc.font = XFont(name=XFONT_NAME, size=10)
            pc.number_format = '€ #,##0.00'

        totaal_cell = ws.cell(row=row, column=5)
        if regel.aantal is not None and regel.prijs_per_stuk is not None:
            col_c = get_column_letter(3)
            col_d = get_column_letter(4)
            totaal_cell.value = f"={col_c}{row}*{col_d}{row}"
        else:
            totaal_cell.value = regel.totaal
        totaal_cell.font = XFont(name=XFONT_NAME, size=10, bold=True)
        totaal_cell.number_format = '€ #,##0.00'

        row += 1

        if regel.foto_base64:
            try:
                foto_bytes = base64.b64decode(regel.foto_base64)
                foto_bytes = _verwerk_foto_bytes(foto_bytes, item_naam=regel.item, max_breedte_px=400)
                img = XLImage(_io.BytesIO(foto_bytes))
                img.width = 90
                img.height = 90
                ws.add_image(img, f"A{row}")
                ws.row_dimensions[row].height = 70
                row += 1
            except Exception as e:
                print(f"[foto-fout] Kon foto voor '{regel.item}' niet in het Excel-bestand plaatsen: {type(e).__name__}: {e}")
    last_item_row = row - 1

    toon_breakdown = (doc_type == "offerte") and bool(data.toeslag_percentage)

    subtotaal_row = None
    if toon_breakdown:
        row += 1
        subtotaal_row = row
        ws.cell(row=row, column=4, value="Subtotaal").font = XFont(name=XFONT_NAME, size=10)
        ws.cell(row=row, column=4).alignment = Alignment(horizontal="right")
        c = ws.cell(row=row, column=5, value=f"=SUM(E{first_item_row}:E{last_item_row})")
        c.font = XFont(name=XFONT_NAME, size=10, bold=True)
        c.number_format = '€ #,##0.00'
        row += 1

    excl_btw_row = None
    if toon_breakdown:
        toeslag_row = row
        ws.cell(row=row, column=4, value=f"Toeslag ({data.toeslag_percentage:g}%)").alignment = Alignment(horizontal="right")
        ws.cell(row=row, column=4).font = XFont(name=XFONT_NAME, size=10)
        c = ws.cell(row=row, column=5, value=f"=E{subtotaal_row}*{data.toeslag_percentage / 100}")
        c.font = XFont(name=XFONT_NAME, size=10)
        c.number_format = '€ #,##0.00'
        row += 1

        excl_btw_row = row
        ws.cell(row=row, column=4, value="Totaal excl. BTW").alignment = Alignment(horizontal="right")
        ws.cell(row=row, column=4).font = XFont(name=XFONT_NAME, size=10)
        c = ws.cell(row=row, column=5, value=f"=E{subtotaal_row}+E{toeslag_row}")
        c.font = XFont(name=XFONT_NAME, size=10)
        c.number_format = '€ #,##0.00'
        row += 1
    elif data.toeslag_percentage:
        # Toeslag blijft verwerkt in het totaal, maar wordt niet als aparte regel
        # (en niet als losse Subtotaal) getoond aan de klant bij orderbevestiging/factuur.
        row += 1
        excl_btw_row = row
        ws.cell(row=row, column=4, value="Totaal excl. BTW").alignment = Alignment(horizontal="right")
        ws.cell(row=row, column=4).font = XFont(name=XFONT_NAME, size=10)
        c = ws.cell(row=row, column=5, value=f"=SUM(E{first_item_row}:E{last_item_row})*(1+{data.toeslag_percentage / 100})")
        c.font = XFont(name=XFONT_NAME, size=10)
        c.number_format = '€ #,##0.00'
        row += 1

    basis_row = excl_btw_row or subtotaal_row
    if basis_row is None:
        # Geen toeslag: basis voor eventuele BTW-berekening is direct de som van de regels.
        row += 1
        subtotaal_row = row
        ws.cell(row=row, column=4, value="Subtotaal").font = XFont(name=XFONT_NAME, size=10)
        ws.cell(row=row, column=4).alignment = Alignment(horizontal="right")
        c = ws.cell(row=row, column=5, value=f"=SUM(E{first_item_row}:E{last_item_row})")
        c.font = XFont(name=XFONT_NAME, size=10, bold=True)
        c.number_format = '€ #,##0.00'
        row += 1
        basis_row = subtotaal_row

    if data.korting_percentage:
        # Korting wordt, net als Toeslag, nergens als aparte regel getoond -
        # alleen stilzwijgend verwerkt door de bestaande basisregel te
        # vermenigvuldigen met de kortingsfactor, zonder een nieuwe rij.
        basis_cell = ws.cell(row=basis_row, column=5)
        bestaande_formule = basis_cell.value
        basis_cell.value = f"=({bestaande_formule[1:]})*(1-{data.korting_percentage / 100})"
        # Het label mag na verwerking van de korting niet meer "Subtotaal"
        # heten, want dat bedrag klopt dan niet meer met de losse regeltotalen.
        label_cell = ws.cell(row=basis_row, column=4)
        if label_cell.value == "Subtotaal":
            label_cell.value = "Totaal excl. BTW" if data.btw_percentage else "TOTAAL"

    if data.btw_percentage:
        btw_row = row
        ws.cell(row=row, column=4, value=f"BTW ({data.btw_percentage:g}%)").alignment = Alignment(horizontal="right")
        ws.cell(row=row, column=4).font = XFont(name=XFONT_NAME, size=10)
        c = ws.cell(row=row, column=5, value=f"=E{basis_row}*{data.btw_percentage / 100}")
        c.font = XFont(name=XFONT_NAME, size=10)
        c.number_format = '€ #,##0.00'
        row += 1

        eind_row = row
        ws.cell(row=row, column=4, value="Totaal incl. BTW").alignment = Alignment(horizontal="right")
        ws.cell(row=row, column=4).font = XFont(name=XFONT_NAME, size=11, bold=True)
        c = ws.cell(row=row, column=5, value=f"=E{basis_row}+E{btw_row}")
        c.font = XFont(name=XFONT_NAME, size=11, bold=True)
        c.number_format = '€ #,##0.00'
        for col in (4, 5):
            ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor=XTOTAL_FILL)
        row += 1
    else:
        eind_row = row
        ws.cell(row=row, column=4, value="TOTAAL").alignment = Alignment(horizontal="right")
        ws.cell(row=row, column=4).font = XFont(name=XFONT_NAME, size=11, bold=True)
        c = ws.cell(row=row, column=5, value=f"=E{basis_row}")
        c.font = XFont(name=XFONT_NAME, size=11, bold=True)
        c.number_format = '€ #,##0.00'
        for col in (4, 5):
            ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor=XTOTAL_FILL)
        row += 1

    if doc_type == "offerte" and data.aanbetaling_termijnen:
        row += 1
        c = ws.cell(row=row, column=1, value="Betaaltermijnen")
        c.font = XFont(name=XFONT_NAME, size=11, bold=True)
        row += 1
        termijnen = data.aanbetaling_termijnen
        term_rows = []
        for idx, termijn in enumerate(termijnen):
            ws.cell(row=row, column=1, value=f"{termijn.label} ({termijn.percentage:g}%)").font = XFont(name=XFONT_NAME, size=10)
            if idx == len(termijnen) - 1 and term_rows:
                refs = "+".join(f"B{r}" for r in term_rows)
                c = ws.cell(row=row, column=2, value=f"=E{eind_row}-({refs})")
            else:
                c = ws.cell(row=row, column=2, value=f"=E{eind_row}*{termijn.percentage / 100}")
            c.font = XFont(name=XFONT_NAME, size=10)
            c.number_format = '€ #,##0.00'
            term_rows.append(row)
            row += 1

    # Let op: algemene_opmerkingen worden bewust NIET in het gegenereerde
    # document geschreven. Dit veld blijft alleen zichtbaar in Lovable zelf
    # (via de API-respons van de parse-stap), niet in wat de klant ontvangt.

    buf = _io.BytesIO()
    wb.save(buf)
    return _fix_openpyxl_lege_formule_waarde(buf.getvalue())


@app.post("/generate-xlsx")
async def generate_xlsx(data: GenerateRequest):
    xlsx_bytes = build_xlsx(data)
    doc_type = data.document_type or "offerte"
    return StreamingResponse(
        _io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{doc_type}.xlsx"'},
    )
